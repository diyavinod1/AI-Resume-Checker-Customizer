"""
Resume Generator Module

Generates formatted resume output in multiple formats (PDF, DOCX, TXT)
with customizable styling, tone, and regional formatting.
"""

import tempfile
import os
from typing import Dict, List, Optional
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

class ResumeGenerator:
    def __init__(self):
        """Initialize the resume generator."""
        self.regional_formats = {
            'US': {
                'date_format': '%m/%d/%Y',
                'phone_format': '(xxx) xxx-xxxx',
                'address_order': ['city', 'state', 'country']
            },
            'UK': {
                'date_format': '%d/%m/%Y', 
                'phone_format': '+44 xxx xxxx xxxx',
                'address_order': ['city', 'country', 'postcode']
            },
            'India': {
                'date_format': '%d-%m-%Y',
                'phone_format': '+91 xxxxx xxxxx',
                'address_order': ['city', 'state', 'country', 'pincode']
            }
        }
    
    def generate_text_resume(self, resume_data: Dict, tone: str = 'professional', region: str = 'US') -> str:
        """Generate a formatted text version of the resume."""
        contact_info = resume_data.get('contact_info', {})
        experience = resume_data.get('experience', [])
        skills = resume_data.get('skills', {})
        education = resume_data.get('education', [])
        projects = resume_data.get('projects', [])
        
        # Build resume text
        resume_lines = []
        
        # Header
        name = contact_info.get('name', 'Professional Resume')
        resume_lines.append(name.upper())
        resume_lines.append('=' * len(name))
        resume_lines.append('')
        
        # Contact Information
        contact_line = []
        if contact_info.get('email'):
            contact_line.append(contact_info['email'])
        if contact_info.get('phone'):
            contact_line.append(self._format_phone(contact_info['phone'], region))
        if contact_info.get('linkedin'):
            contact_line.append(contact_info['linkedin'])
        
        if contact_line:
            resume_lines.append(' | '.join(contact_line))
            resume_lines.append('')
        
        # Professional Summary (if available)
        metadata = resume_data.get('metadata', {})
        if metadata.get('target_role'):
            summary_tone = 'Experienced' if tone == 'professional' else 'Passionate'
            resume_lines.append(f"{summary_tone} {metadata['target_role']} seeking new opportunities.")
            resume_lines.append('')
        
        # Experience Section
        if experience:
            resume_lines.append('PROFESSIONAL EXPERIENCE')
            resume_lines.append('-' * 25)
            resume_lines.append('')
            
            for exp in experience:
                position = exp.get('position', 'Position')
                company = exp.get('company', 'Company')
                duration = exp.get('duration', 'Duration')
                
                # Header line
                resume_lines.append(f"{position} | {company}")
                if duration != 'Not specified':
                    resume_lines.append(f"{duration}")
                resume_lines.append('')
                
                # Description
                description = exp.get('description', [])
                if isinstance(description, str):
                    description = [description]
                
                for desc in description:
                    if desc and desc != 'Not specified':
                        # Apply tone adjustments
                        formatted_desc = self._apply_tone(desc, tone)
                        resume_lines.append(f"• {formatted_desc}")
                
                resume_lines.append('')
        
        # Skills Section
        if skills:
            resume_lines.append('TECHNICAL SKILLS')
            resume_lines.append('-' * 16)
            resume_lines.append('')
            
            for category, skill_list in skills.items():
                if skill_list:
                    category_name = category.replace('_', ' ').title()
                    resume_lines.append(f"{category_name}: {', '.join(skill_list)}")
            
            resume_lines.append('')
        
        # Education Section
        if education:
            resume_lines.append('EDUCATION')
            resume_lines.append('-' * 9)
            resume_lines.append('')
            
            for edu in education:
                degree = edu.get('degree', 'Degree')
                institution = edu.get('institution', 'Institution')
                year = edu.get('year', 'Year')
                
                if degree != 'Not specified':
                    resume_lines.append(f"{degree}")
                if institution != 'Not specified':
                    resume_lines.append(f"{institution}")
                    if year != 'Not specified':
                        resume_lines.append(f"Graduated: {year}")
                resume_lines.append('')
        
        # Projects Section
        if projects:
            resume_lines.append('PROJECTS')
            resume_lines.append('-' * 8)
            resume_lines.append('')
            
            for project in projects:
                title = project.get('title', 'Project')
                description = project.get('description', 'Project description')
                
                resume_lines.append(f"{title}")
                resume_lines.append(f"• {self._apply_tone(description, tone)}")
                resume_lines.append('')
        
        return '\n'.join(resume_lines)
    
    def generate_pdf_resume(self, resume_data: Dict, tone: str = 'professional', region: str = 'US') -> bytes:
        """Generate a PDF version of the resume."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            import io
            
            # Create PDF in memory
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
            
            # Define styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12,
                textColor='#333333'
            )
            
            # Build story (content)
            story = []
            contact_info = resume_data.get('contact_info', {})
            
            # Title
            name = contact_info.get('name', 'Professional Resume')
            story.append(Paragraph(name, title_style))
            
            # Contact info
            contact_line = []
            if contact_info.get('email'):
                contact_line.append(contact_info['email'])
            if contact_info.get('phone'):
                contact_line.append(self._format_phone(contact_info['phone'], region))
            
            if contact_line:
                story.append(Paragraph(' | '.join(contact_line), styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Experience
            experience = resume_data.get('experience', [])
            if experience:
                story.append(Paragraph('PROFESSIONAL EXPERIENCE', heading_style))
                
                for exp in experience:
                    position = exp.get('position', 'Position')
                    company = exp.get('company', 'Company')
                    duration = exp.get('duration', 'Duration')
                    
                    # Job header
                    job_header = f"<b>{position}</b> | {company}"
                    if duration != 'Not specified':
                        job_header += f" | {duration}"
                    
                    story.append(Paragraph(job_header, styles['Normal']))
                    
                    # Description
                    description = exp.get('description', [])
                    if isinstance(description, str):
                        description = [description]
                    
                    for desc in description:
                        if desc and desc != 'Not specified':
                            formatted_desc = self._apply_tone(desc, tone)
                            story.append(Paragraph(f"• {formatted_desc}", styles['Normal']))
                    
                    story.append(Spacer(1, 12))
            
            # Skills
            skills = resume_data.get('skills', {})
            if skills:
                story.append(Paragraph('TECHNICAL SKILLS', heading_style))
                
                for category, skill_list in skills.items():
                    if skill_list:
                        category_name = category.replace('_', ' ').title()
                        skills_text = f"<b>{category_name}:</b> {', '.join(skill_list)}"
                        story.append(Paragraph(skills_text, styles['Normal']))
                
                story.append(Spacer(1, 12))
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            logger.warning("ReportLab not available, falling back to text format")
            # Return text content as bytes if PDF generation fails
            text_content = self.generate_text_resume(resume_data, tone, region)
            return text_content.encode('utf-8')
        except Exception as e:
            logger.error(f"Error generating PDF: {e}")
            text_content = self.generate_text_resume(resume_data, tone, region)
            return text_content.encode('utf-8')
    
    def generate_docx_resume(self, resume_data: Dict, tone: str = 'professional', region: str = 'US') -> bytes:
        """Generate a DOCX version of the resume."""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
            
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            contact_info = resume_data.get('contact_info', {})
            
            # Title
            name = contact_info.get('name', 'Professional Resume')
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(name)
            title_run.bold = True
            title_run.font.size = Inches(0.25)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact info
            contact_line = []
            if contact_info.get('email'):
                contact_line.append(contact_info['email'])
            if contact_info.get('phone'):
                contact_line.append(self._format_phone(contact_info['phone'], region))
            
            if contact_line:
                contact_para = doc.add_paragraph(' | '.join(contact_line))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Empty line
            
            # Experience
            experience = resume_data.get('experience', [])
            if experience:
                exp_heading = doc.add_paragraph('PROFESSIONAL EXPERIENCE')
                exp_heading.runs[0].bold = True
                exp_heading.runs[0].font.size = Inches(0.18)
                
                for exp in experience:
                    position = exp.get('position', 'Position')
                    company = exp.get('company', 'Company')
                    duration = exp.get('duration', 'Duration')
                    
                    # Job header
                    job_para = doc.add_paragraph()
                    job_run = job_para.add_run(f"{position} | {company}")
                    job_run.bold = True
                    
                    if duration != 'Not specified':
                        job_para.add_run(f" | {duration}")
                    
                    # Description
                    description = exp.get('description', [])
                    if isinstance(description, str):
                        description = [description]
                    
                    for desc in description:
                        if desc and desc != 'Not specified':
                            formatted_desc = self._apply_tone(desc, tone)
                            bullet_para = doc.add_paragraph(f"• {formatted_desc}")
                            bullet_para.left_indent = Inches(0.25)
                
                doc.add_paragraph()  # Empty line
            
            # Skills
            skills = resume_data.get('skills', {})
            if skills:
                skills_heading = doc.add_paragraph('TECHNICAL SKILLS')
                skills_heading.runs[0].bold = True
                skills_heading.runs[0].font.size = Inches(0.18)
                
                for category, skill_list in skills.items():
                    if skill_list:
                        category_name = category.replace('_', ' ').title()
                        skills_para = doc.add_paragraph()
                        skills_para.add_run(f"{category_name}: ").bold = True
                        skills_para.add_run(', '.join(skill_list))
            
            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            logger.warning("python-docx not available, falling back to text format")
            text_content = self.generate_text_resume(resume_data, tone, region)
            return text_content.encode('utf-8')
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            text_content = self.generate_text_resume(resume_data, tone, region)
            return text_content.encode('utf-8')
    
    def _apply_tone(self, text: str, tone: str) -> str:
        """Apply tone adjustments to text."""
        if tone == 'professional':
            # Make text more formal
            replacements = {
                'worked on': 'developed',
                'helped with': 'contributed to',
                'did': 'executed',
                'made': 'created',
                'used': 'utilized'
            }
        else:  # casual
            # Keep text more natural
            replacements = {
                'utilized': 'used',
                'executed': 'completed',
                'facilitated': 'helped with'
            }
        
        result = text
        for old, new in replacements.items():
            result = result.replace(old, new)
        
        return result
    
    def _format_phone(self, phone: str, region: str) -> str:
        """Format phone number according to regional standards."""
        # Simple phone formatting - would need more sophisticated logic for production
        if region == 'US':
            # Format as (xxx) xxx-xxxx
            digits = ''.join(filter(str.isdigit, phone))
            if len(digits) == 10:
                return f"({digits[:3]}) {digits[3:6]}-{digits[6:]}"
        elif region == 'UK':
            return f"+44 {phone}"
        elif region == 'India':
            return f"+91 {phone}"
        
        return phone  # Return original if formatting fails
    
    def save_resume_file(self, content: bytes, format_type: str, filename: str = None) -> str:
        """Save resume content to a temporary file and return the path."""
        if not filename:
            filename = f"resume.{format_type.lower()}"
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{format_type.lower()}') as tmp_file:
            tmp_file.write(content)
            tmp_file_path = tmp_file.name
        
        return tmp_file_path

# Usage example
if __name__ == "__main__":
    # Test resume generator
    sample_data = {
        'contact_info': {
            'name': 'John Doe',
            'email': 'john.doe@example.com',
            'phone': '5551234567'
        },
        'experience': [
            {
                'position': 'Software Engineer',
                'company': 'TechCorp',
                'duration': '2020-2023',
                'description': [
                    'Developed web applications using Python and React',
                    'Improved system performance by 30%'
                ]
            }
        ],
        'skills': {
            'programming': ['Python', 'JavaScript'],
            'web': ['React', 'Node.js']
        },
        'education': [
            {
                'degree': "Bachelor's in Computer Science",
                'institution': 'University'
            }
        ]
    }
    
    generator = ResumeGenerator()
    text_resume = generator.generate_text_resume(sample_data)
    print("Generated Text Resume:")
    print(text_resume)